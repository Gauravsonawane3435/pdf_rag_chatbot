import os
import pandas as pd
from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader, UnstructuredImageLoader, TextLoader, CSVLoader
from langchain_core.documents import Document as LC_Document
import pytesseract
from PIL import Image
import openpyxl
from services.multimodal_processor import MultiModalProcessor

class DocumentProcessor:
    multimodal_processor = None
    
    @staticmethod
    def set_multimodal_processor(groq_api_key: str = None):
        """Initialize multi-modal processor with optional vision capabilities."""
        DocumentProcessor.multimodal_processor = MultiModalProcessor(groq_api_key)
    
    @staticmethod
    def process_file(file_path, use_multimodal=True):
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return DocumentProcessor._process_pdf(file_path, use_multimodal)
        elif ext in ['.docx', '.doc']:
            return DocumentProcessor._process_docx(file_path)
        elif ext in ['.xlsx', '.xls']:
            return DocumentProcessor._process_excel(file_path)
        elif ext == '.csv':
            return DocumentProcessor._process_csv(file_path)
        elif ext in ['.png', '.jpg', '.jpeg']:
            return DocumentProcessor._process_image(file_path)
        elif ext == '.txt':
            return DocumentProcessor._process_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    @staticmethod
    def _process_pdf(file_path, use_multimodal=True):
        # Use advanced multi-modal processing if available
        if use_multimodal and DocumentProcessor.multimodal_processor:
            try:
                return DocumentProcessor.multimodal_processor.process_pdf_multimodal(file_path)
            except Exception as e:
                print(f"Multi-modal processing failed, falling back to basic: {e}")
        
        # Fallback to basic PDF processing
        loader = PyPDFLoader(file_path)
        return loader.load()

    @staticmethod
    def _process_docx(file_path):
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return [LC_Document(page_content=text, metadata={"source": file_path, "page": 1})]

    @staticmethod
    def _process_excel(file_path):
        df = pd.read_excel(file_path)
        text = df.to_string()
        return [LC_Document(page_content=text, metadata={"source": file_path, "page": 1})]

    @staticmethod
    def _process_csv(file_path):
        loader = CSVLoader(file_path)
        return loader.load()

    @staticmethod
    def _process_image(file_path):
        # Using pytesseract for OCR
        text = pytesseract.image_to_string(Image.open(file_path))
        return [LC_Document(page_content=text, metadata={"source": file_path, "page": 1})]

    @staticmethod
    def _process_txt(file_path):
        loader = TextLoader(file_path)
        return loader.load()
